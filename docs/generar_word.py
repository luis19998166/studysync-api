"""Genera el documento Word de la guía StudySync a partir del Markdown."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = r"d:\repos\studysync-api\docs\StudySync_Guia_Final_IaC_Prisma_Produccion.md"
OUT_FILE = r"d:\repos\studysync-api\docs\StudySync_Guia_Final_IaC_Prisma_Produccion.docx"

# ── Colores UPDS ─────────────────────────────────────────────────────────────
COLOR_TITULO   = RGBColor(0x1A, 0x56, 0xDB)   # azul
COLOR_H2       = RGBColor(0x1E, 0x40, 0xAF)
COLOR_H3       = RGBColor(0x1D, 0x4E, 0xD8)
COLOR_CODE_BG  = RGBColor(0xF1, 0xF5, 0xF9)
COLOR_CODE_FG  = RGBColor(0x0F, 0x17, 0x2A)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None, mono=False):
    run.font.name  = "Courier New" if mono else name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sizes  = {1: 20, 2: 16, 3: 13, 4: 11}
    colors = {1: COLOR_TITULO, 2: COLOR_H2, 3: COLOR_H3, 4: COLOR_H3}
    set_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level))
    return p

def add_code_block(lines):
    """Agrega un bloque de código con fondo gris claro."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
        # Fondo gris
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F1F5F9')
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        set_font(run, size=9, mono=True, color=COLOR_CODE_FG)

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    # Procesar negrita inline **texto**
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        set_font(run, size=10.5, bold=(i % 2 == 1))
    return p

def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        set_font(run, size=9.5, bold=bold)
        if header:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1A56DB')
            cell._tc.get_or_add_tcPr().append(shd)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

# ── Parsear el Markdown ───────────────────────────────────────────────────────
with open(MD_FILE, encoding='utf-8') as f:
    lines = f.readlines()

in_code = False
code_lines = []
in_table = False
table_obj = None
table_cols = 0

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Bloque de código
    if line.startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            add_code_block(code_lines)
            in_code = False
            code_lines = []
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Tabla Markdown
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # Línea separadora (---|---) → ignorar
        if all(set(c.replace('-', '').replace(':', '').strip()) <= set('') for c in cells):
            i += 1
            continue
        if not in_table:
            in_table = True
            table_cols = len(cells)
            table_obj = doc.add_table(rows=0, cols=table_cols)
            table_obj.style = 'Table Grid'
            add_table_row(table_obj, cells, bold=True, header=True)
        else:
            add_table_row(table_obj, cells)
        i += 1
        continue
    else:
        in_table = False
        table_obj = None

    # Encabezados
    if line.startswith('#### '):
        add_heading(line[5:], 4)
    elif line.startswith('### '):
        add_heading(line[4:], 3)
    elif line.startswith('## '):
        add_heading(line[3:], 2)
    elif line.startswith('# '):
        add_heading(line[2:], 1)
    # Separador horizontal
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1A56DB')
        pBdr.append(bottom)
        pPr.append(pBdr)
    # Listas
    elif line.startswith('- ') or line.startswith('* '):
        add_bullet(line[2:], 0)
    elif re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line)
        add_bullet(text, 0)
    # Línea en blanco
    elif line.strip() == '':
        pass
    # Texto normal
    else:
        # Procesar negrita **texto** en líneas normales
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        parts = re.split(r'\*\*(.+?)\*\*', line)
        for j, part in enumerate(parts):
            # Inline code `codigo`
            subparts = re.split(r'`(.+?)`', part)
            for k, sp in enumerate(subparts):
                if sp:
                    run = p.add_run(sp)
                    set_font(run, size=11, bold=(j % 2 == 1),
                             mono=(k % 2 == 1), color=COLOR_CODE_FG if k % 2 == 1 else None)

    i += 1

doc.save(OUT_FILE)
print(f"Documento generado: {OUT_FILE}")
